# -*- coding: utf-8 -*-
"""
Document Content Extraction Service
Supports PDF, Word (docx), Excel (xlsx), Markdown, Images (OCR)
Word format preserves heading levels, bold, italic, lists, images

Optimizations:
1. GPU acceleration: OCR uses thread pool + GPU parallel processing
2. Smart detection: native PDFs extract text directly, scanned pages use OCR
"""
import os
import traceback
import uuid
import re
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from multiprocessing import cpu_count
from threading import Lock
from config import OUTPUT_FOLDER


# ========== PDF Extraction (multiprocess + PyMuPDF primary + OCR fallback) ==========

def _is_garbage_text(text):
    """
    Check if extracted text is garbled/garbage (language-agnostic, applies to both CJK and Latin).

    Core logic:
    - Count meaningful characters (CJK + Latin letters + digits)
    - If meaningful character ratio is too low → garbage
    - If total meaningful characters too few → insufficient content, needs OCR
    - No longer relies on "Chinese character ratio" — pure English won't be misjudged as garbage

    Returns True if OCR fallback is needed.
    """
    if not text or len(text.strip()) < 5:
        return True

    # Count meaningful character categories
    cjk_chars = len(re.findall(r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]', text))
    latin_chars = len(re.findall(r'[a-zA-Z]', text))
    digit_chars = len(re.findall(r'[0-9]', text))
    meaningful = cjk_chars + latin_chars + digit_chars
    total_chars = len(text.strip())

    if total_chars == 0:
        return True

    # Meaningful character ratio (CJK + Latin + digits all count as meaningful)
    meaningful_ratio = meaningful / total_chars

    # Meaningful ratio below 25% → garbage (too many symbols/punctuation)
    if meaningful_ratio < 0.25:
        return True

    # Too few meaningful characters → insufficient content, needs OCR
    if meaningful < 15:
        return True

    # English document: Latin characters dominate, no CJK validation needed
    # As long as meaningful ratio is sufficient, not garbage
    if latin_chars > 0 and cjk_chars == 0:
        # Pure English: Latin >= 30% is valid content
        if latin_chars / total_chars >= 0.30:
            return False

    # Mixed CJK/English or pure CJK: sufficient meaningful chars → valid
    if meaningful >= 20:
        return False

    return True


def _extract_pdfminer_page(file_path, page_num):
    """
    Extract a single page using pdfminer.six (better CJK CID font support).
    Uses PDFPageExtractor to avoid re-parsing the entire PDF each time.
    """
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer
        page_count = 0
        for page_layout in extract_pages(file_path):
            if page_count == page_num:
                texts = []
                for element in page_layout:
                    if isinstance(element, LTTextContainer):
                        t = element.get_text().strip()
                        if t:
                            texts.append(t)
                return "\n".join(texts)
            page_count += 1
    except Exception:
        pass
    return ""


def _extract_pdfminer_full(file_path):
    """
    Extract full PDF text using pdfminer.six in a single pass.
    Best support for CJK CID fonts.
    """
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(file_path)
        if text and text.strip():
            return text.strip()
    except Exception:
        pass
    return ""


def _extract_text_from_page_fitz(page):
    """
    Extract text from a single PyMuPDF page, trying multiple modes for best result.
    """
    text = page.get_text("text")
    if text and len(text.strip()) > 10 and not _is_garbage_text(text):
        return text.strip()

    # blocks mode (more complete, includes tables etc.)
    blocks_text = page.get_text("blocks")
    if blocks_text:
        lines = []
        for block in blocks_text:
            if len(block) >= 5 and block[4].strip():
                lines.append(block[4].strip())
        if lines:
            combined = "\n".join(lines)
            if not _is_garbage_text(combined) and len(combined) > 20:
                return combined

    # dict mode
    d = page.get_text("dict")
    if d and "blocks" in d:
        parts = []
        for block in d["blocks"]:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        t = span.get("text", "").strip()
                        if t:
                            parts.append(t)
        if parts:
            combined = "\n".join(parts)
            if not _is_garbage_text(combined) and len(combined) > 20:
                return combined

    return text.strip() if text and not _is_garbage_text(text) else ""


def _pdf_page_worker(args):
    """
    Subprocess worker: extract text from a single PDF page.
    Strategy: pdfminer.six (best for CJK) -> PyMuPDF multi-mode -> OCR fallback
    Returns (page_num, text, need_ocr)
    """
    file_path, page_num = args
    try:
        import fitz

        # Strategy 1: pdfminer.six (strongest CJK CID font support)
        text = _extract_pdfminer_page(file_path, page_num)
        if text and len(text) > 5 and not _is_garbage_text(text):
            return (page_num, text, False)

        # Strategy 2: PyMuPDF multi-mode
        doc = fitz.open(file_path)
        page = doc[page_num]
        text = _extract_text_from_page_fitz(page)
        doc.close()
        if text and len(text) > 5 and not _is_garbage_text(text):
            return (page_num, text, False)

        # Strategy 3: OCR fallback (scanned pages)
        return (page_num, None, True)

    except Exception:
        return (page_num, None, True)


def _ocr_pdf_page_worker(args):
    """
    OCR a single PDF page (GPU accelerated version).
    Selects OCR model based on detected language.

    args: (file_path, page_num, lang)
    """
    file_path, page_num, lang = args
    tmp_img = None
    try:
        import fitz

        # Convert PDF page to image using PyMuPDF
        doc = fitz.open(file_path)
        page = doc[page_num]
        mat = fitz.Matrix(2, 2)  # 2x zoom for clarity
        pix = page.get_pixmap(matrix=mat)
        tmp_img = f"_tmp_p{page_num}_{os.getpid()}.png"
        pix.save(tmp_img)
        doc.close()

        # Call OCR service with detected language
        from services.ocr_service import ocr_image
        text = ocr_image(tmp_img, lang=lang)
        return (page_num, text)
    except Exception:
        return (page_num, "")
    finally:
        if tmp_img and os.path.exists(tmp_img):
            try: os.remove(tmp_img)
            except Exception: pass


def extract_text_from_pdf(file_path, file_id=None):
    """
    Extract text from PDF (multiprocess acceleration + auto language detection).
    Strategy: pdfminer.six primary -> PyMuPDF supplement -> parallel OCR fallback (scanned pages)
    V2: auto-detect document language, use English OCR model for English documents.
    """
    try:
        import fitz
        from progress_store import set_progress
        from services.ocr_service import detect_language

        # OCR availability check: don't block regular text PDF extraction if OCR fails
        is_gpu = False
        ocr_available = False
        try:
            from services.ocr_service import is_gpu_available
            is_gpu = is_gpu_available()
            ocr_available = True
        except Exception as ocr_err:
            print(f"[OCR] Not available, using text extraction only: {ocr_err}")

        doc = fitz.open(file_path)
        total_pages = len(doc)
        doc.close()
        if total_pages == 0:
            return ""

        # First, full extraction via pdfminer.six
        if file_id is not None:
            set_progress(file_id, 100, 5, "Extracting text...")
        full_text = _extract_pdfminer_full(file_path)

        # Detect document language (for subsequent OCR model selection)
        doc_lang = detect_language(full_text) if full_text else "en"
        lang_label = {"en": "English", "ch": "Chinese", "mixed": "Chinese+English"}.get(doc_lang, "English")
        print(f"[Lang] Detected document language: {lang_label} → using '{doc_lang}' OCR model")

        if full_text and len(full_text.strip()) > 20 and not _is_garbage_text(full_text):
            # pdfminer extraction succeeded with good quality, use directly
            if file_id is not None:
                set_progress(file_id, 100, 80, "Cleaning data...")
            if file_id is not None:
                set_progress(file_id, 100, 100, "Parse complete")
            return full_text

        # pdfminer insufficient (too little content or poor quality), use multiprocess per-page extraction + OCR fallback
        max_workers = min(cpu_count(), total_pages, 8)

        pages_text = {}
        pages_need_ocr = []

        with ProcessPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(_pdf_page_worker, (file_path, i)): i
                for i in range(total_pages)
            }
            for future in as_completed(futures):
                page_num, text, need_ocr = future.result()
                if need_ocr:
                    pages_need_ocr.append(page_num)
                else:
                    pages_text[page_num] = text
                if file_id is not None:
                    done_count = len(pages_text) + len(pages_need_ocr)
                    pct = int(done_count / total_pages * 50)
                    set_progress(file_id, 100, pct, f"Extracting page {done_count}/{total_pages}")

        # Parallel OCR for pages with no text (scanned pages)
        # GPU acceleration: use thread pool instead of process pool, GPU shared resources work better with threads
        if pages_need_ocr and ocr_available:
            ocr_total = len(pages_need_ocr)
            ocr_done = 0
            # More parallel OCR threads in GPU mode
            max_ocr_workers = min(cpu_count() * 2, ocr_total, 16) if is_gpu else min(cpu_count(), ocr_total, 8)
            with ThreadPoolExecutor(max_workers=max_ocr_workers) as executor:
                # Pass detected language to OCR worker threads
                futures = {
                    executor.submit(_ocr_pdf_page_worker, (file_path, p, doc_lang)): p
                    for p in pages_need_ocr
                }
                for future in as_completed(futures):
                    page_num, text = future.result()
                    ocr_done += 1
                    pages_text[page_num] = text if text else ""
                    if file_id is not None:
                        pct = 50 + int(ocr_done / ocr_total * 50)
                        set_progress(file_id, 100, pct, f"OCR ({lang_label}) {ocr_done}/{ocr_total} pages")
        elif pages_need_ocr and not ocr_available:
            # OCR unavailable, these pages cannot be recognized, leave empty
            for p in pages_need_ocr:
                pages_text[p] = pages_text.get(p, "")

        # Concatenate in order
        parts = []
        for i in range(total_pages):
            text = pages_text.get(i, "")
            if text and text.strip():
                parts.append(f"## Page {i+1}\n{text.strip()}")
        result = "\n\n".join(parts) if parts else ""

        if not result.strip() and pages_need_ocr and not ocr_available:
            raise Exception("PDF parse failed: OCR dependencies (paddle/paddleocr) are required for this scanned PDF")

        if file_id is not None:
            set_progress(file_id, 100, 100, "Parse complete")
        return result

    except Exception as e:
        if file_id is not None:
            from progress_store import set_progress
            set_progress(file_id, 100, 100, "Parse failed")
        raise Exception(f"PDF parse failed: {str(e)}")


# ========== Word (docx) Extraction (preserves formatting and images) ==========

def _extract_run_text(run, paragraph_style):
    """
    Extract formatted text from a single run element.
    Returns Markdown-formatted string.
    """
    text = run.text if run.text else ""
    if not text:
        return ""

    bold = getattr(run, "bold", False)
    italic = getattr(run, "italic", False)
    underline = getattr(run, "underline", False)
    strike = getattr(run, "strike", False)

    # Bold + italic combinations
    if bold and italic:
        text = f"***{text}***"
    elif bold:
        text = f"**{text}**"
    elif italic:
        text = f"*{text}*"
    if underline:
        text = f"<u>{text}</u>"

    return text


def extract_text_from_docx(file_path, output_base_dir=None):
    """
    Extract content from Word document, preserving heading levels, bold, italic, lists, images.
    Images saved to outputs directory, referenced with relative paths in Markdown.

    Args:
        file_path: docx file path
        output_base_dir: image save directory (default: outputs/<filename>/images)
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document(file_path)

        # Set image save directory
        if output_base_dir:
            img_dir = output_base_dir
        else:
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            img_dir = os.path.join(OUTPUT_FOLDER, base_name, "images")
        os.makedirs(img_dir, exist_ok=True)

        result_parts = []

        # Extract embedded images
        extracted_images = {}  # rId -> (saved_filename, full_path)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    img_data = rel.target_part.blob
                    img_ext = rel.target_ref.split(".")[-1].lower()
                    if img_ext not in ["png", "jpg", "jpeg", "gif", "bmp", "webp"]:
                        img_ext = "png"
                    img_name = f"{uuid.uuid4().hex[:12]}.{img_ext}"
                    img_path = os.path.join(img_dir, img_name)
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    extracted_images[rel.rId] = (img_name, img_path.replace("\\", "/"))
                except Exception as img_err:
                    print(f"[docx image extraction failed] {img_err}")

        # Extract body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ""

            # Detect heading levels
            if "heading 1" in style_name or "heading1" in style_name:
                result_parts.append(f"# {text}")
            elif "heading 2" in style_name or "heading2" in style_name:
                result_parts.append(f"## {text}")
            elif "heading 3" in style_name or "heading3" in style_name:
                result_parts.append(f"### {text}")
            elif "heading 4" in style_name or "heading4" in style_name:
                result_parts.append(f"#### {text}")
            elif "title" in style_name:
                result_parts.append(f"# {text}")
            # List items
            elif para.style and ("List" in para.style.name or "Numbering" in para.style.name):
                num_pr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if num_pr is not None:
                    result_parts.append(f"- {text}")
                else:
                    result_parts.append(f"- {text}")
            else:
                # Regular paragraph, iterate runs to preserve formatting
                runs_text = []
                for run in para.runs:
                    run_text = run.text if run.text else ""
                    if not run_text:
                        continue
                    bold = getattr(run, "bold", False)
                    italic = getattr(run, "italic", False)
                    underline = getattr(run, "underline", False)

                    if bold and italic:
                        run_text = f"***{run_text}***"
                    elif bold:
                        run_text = f"**{run_text}**"
                    elif italic:
                        run_text = f"*{run_text}*"

                    runs_text.append(run_text)

                para_text = "".join(runs_text) if runs_text else text
                if para_text.strip():
                    result_parts.append(para_text)

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    if row_idx == 0:
                        table_rows.append("| " + " | ".join(cells) + " |")
                        table_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    else:
                        table_rows.append("| " + " | ".join(cells) + " |")
            if table_rows:
                result_parts.append("\n".join(table_rows))

        return "\n\n".join(result_parts)

    except Exception as e:
        traceback.print_exc()
        raise Exception(f"Word parse failed: {str(e)}")


# ========== Excel Extraction ==========

def extract_text_from_xlsx(file_path):
    """Excel extraction, output as Markdown tables."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(file_path, data_only=True)
        all_sheets = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_data = []
            for row in ws.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_cells = [str(cell) if cell is not None else "" for cell in row]
                    rows_data.append(row_cells)

            if rows_data:
                lines = [f"### Sheet: {sheet_name}\n"]
                for i, row in enumerate(rows_data):
                    if i == 0:
                        lines.append("| " + " | ".join(row) + " |")
                        lines.append("| " + " | ".join(["---"] * len(row)) + " |")
                    else:
                        lines.append("| " + " | ".join(row) + " |")
                all_sheets.append("\n".join(lines))

        return "\n\n".join(all_sheets)

    except Exception as e:
        traceback.print_exc()
        raise Exception(f"Excel parse failed: {str(e)}")


# ========== Markdown Read ==========

def extract_text_from_markdown(file_path):
    """Read Markdown file content."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        raise Exception(f"Markdown file read failed: {str(e)}")


# ========== Image OCR ==========

def extract_text_from_image(file_path, lang="en"):
    """
    OCR text recognition for images.

    Args:
        file_path: image file path
        lang: OCR language model — "en" (English, default), "ch" (Chinese)
    """
    from services.ocr_service import ocr_image
    return ocr_image(file_path, lang=lang)


# ========== Unified Entry Point ==========

def extract_text(file_path, file_ext, file_id=None):
    """
    Auto-select extraction method based on file extension.

    Args:
        file_path: absolute file path
        file_ext: file extension (lowercase, without dot)
        file_id: file ID for progress reporting (PDF only supports detailed progress)
    """
    extractors = {
        "pdf": lambda fp: extract_text_from_pdf(fp, file_id=file_id),
        "docx": lambda fp: _docx_with_progress(fp, file_id),
        "xlsx": lambda fp: _xlsx_with_progress(fp, file_id),
        "md": extract_text_from_markdown,
        "png": lambda fp: _image_with_progress(fp, file_id),
        "jpg": lambda fp: _image_with_progress(fp, file_id),
        "jpeg": lambda fp: _image_with_progress(fp, file_id),
        "jfif": lambda fp: _image_with_progress(fp, file_id),
        "gif": lambda fp: _image_with_progress(fp, file_id),
        "bmp": lambda fp: _image_with_progress(fp, file_id),
        "webp": lambda fp: _image_with_progress(fp, file_id),
    }

    extractor = extractors.get(file_ext.lower())
    if not extractor:
        raise Exception(f"Unsupported file format: .{file_ext}")

    return extractor(file_path)


def _docx_with_progress(file_path, file_id):
    """docx extraction with progress reporting"""
    from progress_store import set_progress
    if file_id is not None:
        set_progress(file_id, 100, 30, "Parsing Word document...")
    result = extract_text_from_docx(file_path)
    if file_id is not None:
        set_progress(file_id, 100, 80, "Done")
    return result


def _xlsx_with_progress(file_path, file_id):
    """xlsx extraction with progress reporting"""
    from progress_store import set_progress
    if file_id is not None:
        set_progress(file_id, 100, 50, "Parsing Excel...")
    result = extract_text_from_xlsx(file_path)
    if file_id is not None:
        set_progress(file_id, 100, 90, "Done")
    return result


def _image_with_progress(file_path, file_id):
    """Image OCR with progress reporting (uses configured default language)"""
    from progress_store import set_progress
    from config import OCR_LANG
    if file_id is not None:
        lang_label = "English" if OCR_LANG == "en" else "Chinese"
        set_progress(file_id, 100, 20, f"OCR in progress ({lang_label})...")
    result = extract_text_from_image(file_path, lang=OCR_LANG)
    if file_id is not None:
        set_progress(file_id, 100, 90, "OCR complete")
    return result
